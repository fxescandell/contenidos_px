import json as _json
import platform
import os
from datetime import datetime
from fastapi import APIRouter, Depends, Request, Form, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
from typing import List, Dict, Any

from app.db.session import get_db, engine, SessionLocal
from app.db.models import SourceBatch
from app.db.settings_models import SystemSetting
from app.services.settings.service import SettingsService, SettingsResolver
from app.schemas.settings import SettingItemUpdate
from app.core.settings_enums import SettingType
from app.services.remote.clients import FtpRemoteInboxClient, SftpRemoteInboxClient, LocalFolderInboxClient, SmbRemoteInboxClient, FtpOutfolderClient, LocalOutfolderClient
from app.services.notifications.telegram import TelegramNotifier
from app.services.categories.service import get_category_export_configs

router = APIRouter()
templates = Jinja2Templates(directory="app/templates")


def _extract_http_error_message(exc: Exception) -> str:
    response = getattr(exc, "response", None)
    if response is None:
        return str(exc)

    try:
        payload = response.json()
    except Exception:
        payload = None

    if isinstance(payload, dict):
        error = payload.get("error")
        if isinstance(error, dict):
            details = error.get("message") or error.get("status") or error.get("code")
            if details:
                return str(details)
        message = payload.get("message")
        if message:
            return str(message)

    text = getattr(response, "text", "") or str(exc)
    return text[:300]


def _format_llm_test_error(provider: str, model: str, exc: Exception) -> str:
    response = getattr(exc, "response", None)
    status_code = getattr(response, "status_code", None)
    base_message = _extract_http_error_message(exc)
    provider_name = provider or "LLM"
    model_name = model or "modelo sin especificar"

    if provider == "gemini":
        if status_code == 429:
            return (
                f"Gemini ({model_name}) ha rechazado la prueba por limite de peticiones o cuota (HTTP 429). "
                "Es habitual en modelos preview o cuando la API key no tiene cuota disponible. "
                "Prueba con un modelo mas estable como gemini-2.0-flash o gemini-2.0-flash-lite y vuelve a intentarlo. "
                f"Detalle: {base_message}"
            )
        if status_code == 404:
            return (
                f"Gemini no encuentra el modelo '{model_name}' o no esta disponible para esta API key. "
                "Carga la lista de modelos de nuevo y selecciona uno soportado por tu cuenta. "
                f"Detalle: {base_message}"
            )
        if status_code == 400:
            return f"Gemini ha rechazado la solicitud para el modelo '{model_name}'. Revisa que el modelo soporte generateContent. Detalle: {base_message}"

    if status_code:
        return f"{provider_name} ({model_name}) devolvio HTTP {status_code}. Detalle: {base_message}"

    return base_message

def _get_category_items(db: Session, category: str):
    SettingsService.initialize_defaults(db)
    return SettingsService.get_section(db, category)

@router.get("/", response_class=HTMLResponse)
def settings_dashboard(request: Request, db: Session = Depends(get_db)):
    """Panel principal de configuración"""
    SettingsService.initialize_defaults(db)
    return templates.TemplateResponse(
        request=request,
        name="settings/dashboard.html",
        context={"status": "loaded", "active_page": "general"}
    )

@router.get("/general", response_class=HTMLResponse)
def settings_general(request: Request, db: Session = Depends(get_db)):
    SettingsService.initialize_defaults(db)
    SettingsResolver.reload(db)
    items = SettingsService.get_section(db, "general")
    values = {}
    for item in items:
        values[item["key"]] = item["value"]
    return templates.TemplateResponse(
        request=request,
        name="settings/general.html",
        context={"values": values, "active_page": "general", "success": request.query_params.get("success")}
    )

@router.post("/general")
async def save_general(request: Request, db: Session = Depends(get_db)):
    form_data = await request.form()
    updates = []
    type_map = {
        "project_name": SettingType.STRING,
        "app_env": SettingType.STRING,
        "enable_auto_processing": SettingType.BOOLEAN,
        "enable_watcher": SettingType.BOOLEAN,
        "watcher_interval_seconds": SettingType.INTEGER,
        "min_file_size_kb": SettingType.INTEGER,
        "enable_telegram_on_error": SettingType.BOOLEAN,
        "max_parallel_batches": SettingType.INTEGER,
    }
    for key, val_type in type_map.items():
        val = form_data.get(f"setting_{key}")
        is_secret = form_data.get(f"secret_{key}") == "on"
        final_val = val
        if val_type == SettingType.BOOLEAN:
            final_val = val == "true"
        elif val_type == SettingType.INTEGER:
            try:
                final_val = int(val) if val else 0
            except (ValueError, TypeError):
                final_val = 0
        updates.append(SettingItemUpdate(key=key, value=final_val, value_type=val_type, is_secret=is_secret))
    SettingsService.update_section(db, "general", updates, user="admin_user")
    return RedirectResponse(url="/settings/general?success=true", status_code=303)

@router.get("/system-info")
def system_info(db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    db_info = {}
    try:
        from sqlalchemy import text
        db_info["engine"] = str(engine.url)
        with SessionLocal() as session:
            from app.db.models import SourceBatch
            count = session.query(SourceBatch).count()
            db_info["batches_total"] = count
            db_info["db_size"] = f"{os.path.getsize('editorial.db') / 1024:.1f} KB" if os.path.exists("editorial.db") else "N/A"
    except Exception as e:
        db_info["error"] = str(e)

    import psutil
    mem = psutil.virtual_memory()
    cpu_pct = psutil.cpu_percent(interval=1)

    return {
        "python_version": platform.python_version(),
        "platform": platform.platform(),
        "hostname": platform.node(),
        "app_env": SettingsResolver.get("app_env", "production"),
        "project_name": SettingsResolver.get("project_name", "Editorial WP"),
        "auto_processing": SettingsResolver.get("enable_auto_processing", True),
        "watcher_enabled": SettingsResolver.get("enable_watcher", True),
        "database": db_info,
        "memory_used_mb": round(mem.used / 1024 / 1024, 1),
        "memory_total_mb": round(mem.total / 1024 / 1024, 1),
        "memory_percent": mem.percent,
        "cpu_percent": cpu_pct,
        "uptime": str(datetime.now()),
    }

@router.get("/test/db-health")
def test_db_health(db: Session = Depends(get_db)):
    results = []
    try:
        from sqlalchemy import text
        db.execute(text("SELECT 1"))
        results.append("SELECT 1: OK")
    except Exception as e:
        results.append(f"SELECT 1: ERROR - {e}")

    try:
        from app.db.models import SourceBatch, CanonicalContent, ProcessingEvent
        from app.db.settings_models import SystemSetting
        batch_count = db.query(SourceBatch).count()
        setting_count = db.query(SystemSetting).count()
        results.append(f"SourceBatch: {batch_count} registros")
        results.append(f"SystemSetting: {setting_count} registros")
    except Exception as e:
        results.append(f"Modelos: ERROR - {e}")

    try:
        from app.db.models import CanonicalContent, ProcessingEvent
        cc = db.query(CanonicalContent).count()
        ev = db.query(ProcessingEvent).count()
        results.append(f"CanonicalContent: {cc} registros")
        results.append(f"ProcessingEvent: {ev} registros")
    except Exception as e:
        results.append(f"Modelos extra: ERROR - {e}")

    ok = "ERROR" not in " ".join(results)
    return {"success": ok, "message": " | ".join(results)}

@router.get("/inbox", response_class=HTMLResponse)
def settings_inbox(request: Request, db: Session = Depends(get_db)):
    SettingsService.initialize_defaults(db)
    SettingsResolver.reload(db)
    items = SettingsService.get_section(db, "inbox")
    values = {}
    for item in items:
        values[item["key"]] = item["value"]
    current_mode = values.get("hot_folder_mode", "local")
    if current_mode == "********":
        current_mode = "local"

    folders_json = values.get("hotfolder_folders", "[]")
    try:
        folders = _json.loads(folders_json) if isinstance(folders_json, str) else folders_json
    except Exception:
        folders = []
    if not isinstance(folders, list):
        folders = []
    local_folders_json = values.get("hotfolder_local_folders", "[]")
    try:
        local_folders = _json.loads(local_folders_json) if isinstance(local_folders_json, str) else local_folders_json
    except Exception:
        local_folders = []
    if not isinstance(local_folders, list):
        local_folders = []

    return templates.TemplateResponse(
        request=request,
        name="settings/inbox.html",
        context={
            "active_page": "inbox",
            "active_mode": SettingsResolver.get("active_source_mode", "smb") or "smb",
            "current_mode": current_mode,
            "values": values,
            "folders": folders,
            "folders_json": _json.dumps(folders),
            "local_folders": local_folders,
            "local_folders_json": _json.dumps(local_folders),
            "success": request.query_params.get("success")
        }
    )

@router.post("/inbox")
async def save_inbox(request: Request, db: Session = Depends(get_db)):
    form_data = await request.form()
    updates = []
    type_map = {
        "hot_folder_mode": SettingType.STRING,
        "hot_folder_local_path": SettingType.STRING,
        "remote_inbox_host": SettingType.STRING,
        "remote_inbox_port": SettingType.INTEGER,
        "remote_inbox_username": SettingType.STRING,
        "remote_inbox_password": SettingType.STRING,
        "remote_inbox_base_path": SettingType.STRING,
        "remote_inbox_passive_mode": SettingType.BOOLEAN,
        "remote_inbox_processed_path": SettingType.STRING,
        "remote_inbox_timeout": SettingType.INTEGER,
        "smb_share_name": SettingType.STRING,
        "smb_domain": SettingType.STRING,
    }
    for key, val_type in type_map.items():
        val = form_data.get(f"setting_{key}")
        is_secret = form_data.get(f"secret_{key}") == "on"
        final_val = val
        if val_type == SettingType.BOOLEAN:
            final_val = val == "true"
        elif val_type == SettingType.INTEGER:
            try:
                final_val = int(val) if val else 0
            except (ValueError, TypeError):
                final_val = 0
        updates.append(SettingItemUpdate(
            key=key,
            value=final_val,
            value_type=val_type,
            is_secret=is_secret
        ))

    folders_raw = form_data.get("hotfolder_folders_json", "[]")
    updates.append(SettingItemUpdate(
        key="hotfolder_folders",
        value=folders_raw,
        value_type=SettingType.STRING,
        is_secret=False
    ))
    local_folders_raw = form_data.get("hotfolder_local_folders_json", "[]")
    updates.append(SettingItemUpdate(
        key="hotfolder_local_folders",
        value=local_folders_raw,
        value_type=SettingType.STRING,
        is_secret=False
    ))

    SettingsService.update_section(db, "inbox", updates, user="admin_user")
    return RedirectResponse(url="/settings/inbox?success=true", status_code=303)

@router.get("/outfolder", response_class=HTMLResponse)
def settings_outfolder(request: Request, db: Session = Depends(get_db)):
    SettingsService.initialize_defaults(db)
    SettingsResolver.reload(db)
    items = SettingsService.get_section(db, "outfolder")
    values = {}
    for item in items:
        values[item["key"]] = item["value"]
    current_mode = values.get("outfolder_mode", "ftp")
    if current_mode == "********":
        current_mode = "ftp"
    folders_json = values.get("outfolder_folders", "[]")
    try:
        folders = _json.loads(folders_json) if isinstance(folders_json, str) else folders_json
    except Exception:
        folders = []
    if not isinstance(folders, list):
        folders = []
    local_folders_json = values.get("outfolder_local_folders", "[]")
    try:
        local_folders = _json.loads(local_folders_json) if isinstance(local_folders_json, str) else local_folders_json
    except Exception:
        local_folders = []
    if not isinstance(local_folders, list):
        local_folders = []
    return templates.TemplateResponse(
        request=request,
        name="settings/outfolder.html",
        context={
            "active_page": "outfolder",
            "active_mode": SettingsResolver.get("active_source_mode", "smb") or "smb",
            "current_mode": current_mode,
            "values": values,
            "folders": folders,
            "folders_json": _json.dumps(folders),
            "local_folders": local_folders,
            "local_folders_json": _json.dumps(local_folders),
            "success": request.query_params.get("success")
        }
    )

@router.post("/outfolder")
async def save_outfolder(request: Request, db: Session = Depends(get_db)):
    form_data = await request.form()
    updates = []
    type_map = {
        "outfolder_mode": SettingType.STRING,
        "outfolder_local_path": SettingType.STRING,
        "outfolder_host": SettingType.STRING,
        "outfolder_port": SettingType.INTEGER,
        "outfolder_username": SettingType.STRING,
        "outfolder_password": SettingType.STRING,
        "outfolder_timeout": SettingType.INTEGER,
        "outfolder_passive_mode": SettingType.BOOLEAN,
        "outfolder_public_base_url": SettingType.STRING,
    }
    for key, val_type in type_map.items():
        val = form_data.get(f"setting_{key}")
        is_secret = form_data.get(f"secret_{key}") == "on"
        final_val = val
        if val_type == SettingType.BOOLEAN:
            final_val = val == "true"
        elif val_type == SettingType.INTEGER:
            try:
                final_val = int(val) if val else 0
            except (ValueError, TypeError):
                final_val = 0
        updates.append(SettingItemUpdate(key=key, value=final_val, value_type=val_type, is_secret=is_secret))
    folders_raw = form_data.get("outfolder_folders_json", "[]")
    updates.append(SettingItemUpdate(key="outfolder_folders", value=folders_raw, value_type=SettingType.STRING, is_secret=False))
    local_folders_raw = form_data.get("outfolder_local_folders_json", "[]")
    updates.append(SettingItemUpdate(key="outfolder_local_folders", value=local_folders_raw, value_type=SettingType.STRING, is_secret=False))
    SettingsService.update_section(db, "outfolder", updates, user="admin_user")
    return RedirectResponse(url="/settings/outfolder?success=true", status_code=303)

# --- Dedicated Routes (before catch-all /{category}) ---

@router.get("/telegram", response_class=HTMLResponse)
def settings_telegram(request: Request, db: Session = Depends(get_db)):
    SettingsService.initialize_defaults(db)
    SettingsResolver.reload(db)
    items = SettingsService.get_section(db, "telegram")
    values = {}
    for item in items:
        values[item["key"]] = item["value"]
    return templates.TemplateResponse(
        request=request,
        name="settings/telegram.html",
        context={"values": values, "active_page": "telegram", "success": request.query_params.get("success")}
    )

@router.post("/telegram")
async def save_telegram(request: Request, db: Session = Depends(get_db)):
    form_data = await request.form()
    updates = []
    type_map = {
        "telegram_enabled": SettingType.BOOLEAN,
        "telegram_bot_token": SettingType.STRING,
        "telegram_chat_id": SettingType.STRING,
    }
    for key, val_type in type_map.items():
        val = form_data.get(f"setting_{key}")
        is_secret = form_data.get(f"secret_{key}") == "on"
        final_val = val
        if val_type == SettingType.BOOLEAN:
            final_val = val == "true"
        updates.append(SettingItemUpdate(key=key, value=final_val, value_type=val_type, is_secret=is_secret))
    SettingsService.update_section(db, "telegram", updates, user="admin_user")
    return RedirectResponse(url="/settings/telegram?success=true", status_code=303)

@router.get("/processing", response_class=HTMLResponse)
def settings_processing(request: Request, db: Session = Depends(get_db)):
    SettingsService.initialize_defaults(db)
    SettingsResolver.reload(db)
    items = SettingsService.get_section(db, "processing")
    values = {}
    for item in items:
        values[item["key"]] = item["value"]
    return templates.TemplateResponse(
        request=request,
        name="settings/processing.html",
        context={"values": values, "active_page": "processing", "success": request.query_params.get("success")}
    )

@router.post("/processing")
async def save_processing(request: Request, db: Session = Depends(get_db)):
    form_data = await request.form()
    updates = []
    type_map = {
        "auto_publish_enabled": SettingType.BOOLEAN,
        "enable_retry_on_failure": SettingType.BOOLEAN,
        "scan_interval_seconds": SettingType.INTEGER,
        "batch_size_limit": SettingType.INTEGER,
        "skip_duplicate_files": SettingType.BOOLEAN,
        "enable_ocr": SettingType.BOOLEAN,
    }
    for key, val_type in type_map.items():
        val = form_data.get(f"setting_{key}")
        is_secret = form_data.get(f"secret_{key}") == "on"
        final_val = val
        if val_type == SettingType.BOOLEAN:
            final_val = val == "true"
        elif val_type == SettingType.INTEGER:
            try:
                final_val = int(val) if val else 0
            except (ValueError, TypeError):
                final_val = 0
        updates.append(SettingItemUpdate(key=key, value=final_val, value_type=val_type, is_secret=is_secret))
    SettingsService.update_section(db, "processing", updates, user="admin_user")
    return RedirectResponse(url="/settings/processing?success=true", status_code=303)

@router.get("/publishing", response_class=HTMLResponse)
def settings_publishing(request: Request, db: Session = Depends(get_db)):
    SettingsService.initialize_defaults(db)
    SettingsResolver.reload(db)
    items = SettingsService.get_section(db, "publishing")
    values = {}
    for item in items:
        values[item["key"]] = item["value"]
    return templates.TemplateResponse(
        request=request,
        name="settings/publishing.html",
        context={"values": values, "active_page": "publishing", "success": request.query_params.get("success")}
    )

@router.post("/publishing")
async def save_publishing(request: Request, db: Session = Depends(get_db)):
    form_data = await request.form()
    updates = []
    type_map = {
        "export_mode": SettingType.STRING,
        "export_include_media": SettingType.BOOLEAN,
        "wp_api_url": SettingType.STRING,
        "wp_username": SettingType.STRING,
        "wp_app_password": SettingType.STRING,
        "wp_default_status": SettingType.STRING,
        "wp_default_author_id": SettingType.INTEGER,
    }
    for key, val_type in type_map.items():
        val = form_data.get(f"setting_{key}")
        is_secret = form_data.get(f"secret_{key}") == "on"
        final_val = val
        if val_type == SettingType.BOOLEAN:
            final_val = val == "true"
        elif val_type == SettingType.INTEGER:
            try:
                final_val = int(val) if val else 0
            except (ValueError, TypeError):
                final_val = 0
        updates.append(SettingItemUpdate(key=key, value=final_val, value_type=val_type, is_secret=is_secret))
    SettingsService.update_section(db, "publishing", updates, user="admin_user")
    return RedirectResponse(url="/settings/publishing?success=true", status_code=303)

@router.get("/ai", response_class=HTMLResponse)
def settings_ai(request: Request, db: Session = Depends(get_db)):
    SettingsService.initialize_defaults(db)
    SettingsResolver.reload(db)
    items = SettingsService.get_section(db, "ai")
    values = {}
    for item in items:
        values[item["key"]] = item["value"]

    connections_json = values.get("llm_connections", "[]")
    if connections_json == "********":
        connections_json = "[]"
    try:
        connections = _json.loads(connections_json)
    except Exception:
        connections = []
    if not isinstance(connections, list):
        connections = []

    if not connections:
        provider = values.get("llm_provider", "")
        api_key = values.get("llm_api_key", "")
        if api_key and api_key != "********" and provider:
            from uuid import uuid4
            connections = [{
                "id": str(uuid4()),
                "name": provider.title(),
                "provider": provider,
                "api_key": api_key,
                "model": values.get("llm_model", ""),
                "temperature": float(values.get("llm_temperature", 0.3) or 0.3),
                "enabled": values.get("llm_enabled", False),
                "active": True,
                "models": []
            }]
            connections_json = _json.dumps(connections)

    return templates.TemplateResponse(
        request=request,
        name="settings/ai.html",
        context={
            "values": values,
            "active_page": "ai",
            "connections": connections,
            "connections_json": connections_json,
            "success": request.query_params.get("success")
        }
    )

@router.post("/ai")
async def save_ai(request: Request, db: Session = Depends(get_db)):
    form_data = await request.form()
    updates = []

    connections_json = form_data.get("llm_connections_json", "[]")
    try:
        connections = _json.loads(connections_json)
    except Exception:
        connections = []

    updates.append(SettingItemUpdate(
        key="llm_connections",
        value=connections_json,
        value_type=SettingType.JSON,
        is_secret=False
    ))

    active = next((c for c in connections if c.get("active") and c.get("enabled")), None)
    if active:
        updates.append(SettingItemUpdate(key="llm_enabled", value=True, value_type=SettingType.BOOLEAN, is_secret=False))
        updates.append(SettingItemUpdate(key="llm_provider", value=active.get("provider", ""), value_type=SettingType.STRING, is_secret=False))
        updates.append(SettingItemUpdate(key="llm_api_key", value=active.get("api_key", ""), value_type=SettingType.STRING, is_secret=True))
        updates.append(SettingItemUpdate(key="llm_model", value=active.get("model", ""), value_type=SettingType.STRING, is_secret=False))
        temp = active.get("temperature", 0.3)
        try:
            temp = float(temp)
        except (ValueError, TypeError):
            temp = 0.3
        updates.append(SettingItemUpdate(key="llm_temperature", value=temp, value_type=SettingType.FLOAT, is_secret=False))
    else:
        updates.append(SettingItemUpdate(key="llm_enabled", value=False, value_type=SettingType.BOOLEAN, is_secret=False))

    updates.append(SettingItemUpdate(key="ocr_engine", value=form_data.get("setting_ocr_engine", "disabled"), value_type=SettingType.STRING, is_secret=False))
    updates.append(SettingItemUpdate(key="ocr_language", value=form_data.get("setting_ocr_language", "cat+spa"), value_type=SettingType.STRING, is_secret=False))
    dpi_val = form_data.get("setting_ocr_dpi", "300")
    try:
        dpi_val = int(dpi_val)
    except (ValueError, TypeError):
        dpi_val = 300
    updates.append(SettingItemUpdate(key="ocr_dpi", value=dpi_val, value_type=SettingType.INTEGER, is_secret=False))
    updates.append(SettingItemUpdate(key="ocr_vision_connection_id", value=form_data.get("setting_ocr_vision_connection_id", ""), value_type=SettingType.STRING, is_secret=False))

    SettingsService.update_section(db, "ai", updates, user="admin_user")
    return RedirectResponse(url="/settings/ai?success=true", status_code=303)

@router.get("/categories", response_class=HTMLResponse)
def settings_categories(request: Request, db: Session = Depends(get_db)):
    SettingsService.initialize_defaults(db)
    SettingsResolver.reload(db)
    configs = get_category_export_configs()
    return templates.TemplateResponse(
        request=request,
        name="settings/categories.html",
        context={
            "configs": configs,
            "configs_json": _json.dumps(configs, ensure_ascii=False),
            "active_page": "categories",
            "success": request.query_params.get("success")
        }
    )

@router.post("/categories")
async def save_categories(request: Request, db: Session = Depends(get_db)):
    form_data = await request.form()
    configs_json = form_data.get("category_export_configs_json", "[]")
    try:
        configs = _json.loads(configs_json)
        if not isinstance(configs, list):
            configs = []
    except Exception:
        configs = []

    updates = [
        SettingItemUpdate(
            key="category_export_configs",
            value=_json.dumps(configs, ensure_ascii=False),
            value_type=SettingType.JSON,
            is_secret=False
        )
    ]
    SettingsService.update_section(db, "categories", updates, user="admin_user")
    return RedirectResponse(url="/settings/categories?success=true", status_code=303)

@router.get("/paths", response_class=HTMLResponse)
def settings_paths(request: Request, db: Session = Depends(get_db)):
    SettingsService.initialize_defaults(db)
    SettingsResolver.reload(db)
    items = SettingsService.get_section(db, "paths")
    values = {}
    for item in items:
        values[item["key"]] = item["value"]
    return templates.TemplateResponse(
        request=request,
        name="settings/paths.html",
        context={"values": values, "active_page": "paths", "success": request.query_params.get("success")}
    )

@router.post("/paths")
async def save_paths(request: Request, db: Session = Depends(get_db)):
    form_data = await request.form()
    updates = []
    type_map = {
        "working_folder_path": SettingType.STRING,
        "cleanup_working_folder_after_success": SettingType.BOOLEAN,
        "export_output_path": SettingType.STRING,
        "temp_folder_path": SettingType.STRING,
        "log_folder_path": SettingType.STRING,
    }
    for key, val_type in type_map.items():
        if val_type == SettingType.BOOLEAN:
            val = form_data.get(f"setting_{key}") == "true"
        else:
            val = form_data.get(f"setting_{key}")
        is_secret = form_data.get(f"secret_{key}") == "on"
        updates.append(SettingItemUpdate(key=key, value=val if val_type == SettingType.BOOLEAN else (val or ""), value_type=val_type, is_secret=is_secret))
    SettingsService.update_section(db, "paths", updates, user="admin_user")
    return RedirectResponse(url="/settings/paths?success=true", status_code=303)

@router.get("/tools", response_class=HTMLResponse)
def settings_tools(request: Request):
    return templates.TemplateResponse(
        request=request,
        name="settings/tools.html",
        context={"success": request.query_params.get("success"), "active_page": "tools"}
    )

@router.get("/flows", response_class=HTMLResponse)
def settings_flows(request: Request):
    return templates.TemplateResponse(
        request=request,
        name="settings/flows.html",
        context={"success": request.query_params.get("success"), "active_page": "flows"}
    )

@router.post("/test/run-all")
async def run_all_tests(request: Request, db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    results = {}

    try:
        from sqlalchemy import text
        db.execute(text("SELECT 1"))
        results["database"] = {"success": True, "message": "Conexion OK"}
    except Exception as e:
        results["database"] = {"success": False, "message": str(e)}

    try:
        from sqlalchemy import text
        db.execute(text("SELECT key, category FROM system_settings LIMIT 1"))
        results["settings_table"] = {"success": True, "message": "Settings DB OK"}
    except Exception as e:
        results["settings_table"] = {"success": False, "message": str(e)}

    try:
        from app.db.models import SourceBatch
        count = db.query(SourceBatch).count()
        results["models"] = {"success": True, "message": f"SourceBatch: {count} registros"}
    except Exception as e:
        results["models"] = {"success": False, "message": str(e)}

    try:
        import os as _os
        paths_to_check = {
            "working": SettingsResolver.get("working_folder_path", "/tmp/editorial_working"),
            "export": SettingsResolver.get("export_output_path", "/tmp/editorial_export"),
            "temp": SettingsResolver.get("temp_folder_path", "/tmp/editorial_temp"),
            "logs": SettingsResolver.get("log_folder_path", "logs"),
        }
        path_results = []
        for name, path in paths_to_check.items():
            if not path:
                path_results.append(f"{name}: sin configurar")
                continue
            exists = _os.path.exists(path)
            writable = _os.access(path, _os.W_OK) if exists else False
            status = "OK" if writable else ("sin permisos" if exists else "no existe")
            path_results.append(f"{name}: {status}")
        ok = all("OK" in p for p in path_results)
        results["paths"] = {"success": ok, "message": "; ".join(path_results)}
    except Exception as e:
        results["paths"] = {"success": False, "message": str(e)}

    try:
        import psutil
        mem = psutil.virtual_memory()
        cpu = psutil.cpu_percent(interval=0.5)
        results["system"] = {"success": True, "message": f"CPU: {cpu}%, RAM: {mem.percent}% ({round(mem.used/1024/1024)}/{round(mem.total/1024/1024)} MB)"}
    except Exception as e:
        results["system"] = {"success": False, "message": str(e)}

    try:
        import requests as _req
        conn_json = SettingsResolver.get("llm_connections", "[]")
        connections = _json.loads(conn_json)
        active = next((c for c in connections if c.get("active") and c.get("enabled")), None)
        if not active:
            results["llm"] = {"success": False, "message": "No hay conexion LLM activa"}
        else:
            provider = active.get("provider", "")
            api_key = active.get("api_key", "")
            model = active.get("model", "")
            if provider == "openai":
                url = "https://api.openai.com/v1/chat/completions"
                headers = {"Authorization": f"Bearer {api_key}"}
                payload = {"model": model, "messages": [{"role": "user", "content": "OK"}], "max_tokens": 3}
            elif provider == "anthropic":
                url = "https://api.anthropic.com/v1/messages"
                headers = {"x-api-key": api_key, "anthropic-version": "2023-06-01", "content-type": "application/json"}
                payload = {"model": model, "max_tokens": 3, "messages": [{"role": "user", "content": "OK"}]}
            elif provider == "groq":
                url = "https://api.groq.com/openai/v1/chat/completions"
                headers = {"Authorization": f"Bearer {api_key}"}
                payload = {"model": model, "messages": [{"role": "user", "content": "OK"}], "max_tokens": 3}
            elif provider == "mistral":
                url = "https://api.mistral.ai/v1/chat/completions"
                headers = {"Authorization": f"Bearer {api_key}"}
                payload = {"model": model, "messages": [{"role": "user", "content": "OK"}], "max_tokens": 3}
            elif provider == "gemini":
                url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
                headers = {}
                payload = {"contents": [{"parts": [{"text": "OK"}]}], "generationConfig": {"maxOutputTokens": 3}}
            elif provider == "ollama":
                url = "http://localhost:11434/api/chat"
                headers = {"Content-Type": "application/json"}
                payload = {"model": model, "messages": [{"role": "user", "content": "OK"}], "stream": False}
            else:
                raise Exception(f"Proveedor {provider} no soportado en batch test")
            resp = _req.post(url, json=payload, headers=headers, timeout=15)
            resp.raise_for_status()
            results["llm"] = {"success": True, "message": f"{provider} ({model}) OK"}
    except Exception as e:
        results["llm"] = {"success": False, "message": _format_llm_test_error(provider if 'provider' in locals() else '', model if 'model' in locals() else '', e)}

    try:
        client = SmbRemoteInboxClient()
        success, msg = client.test_connection()
        results["smb"] = {"success": success, "message": msg}
    except Exception as e:
        results["smb"] = {"success": False, "message": str(e)}

    try:
        client = FtpOutfolderClient()
        success, msg = client.test_connection()
        results["outfolder"] = {"success": success, "message": msg}
    except Exception as e:
        results["outfolder"] = {"success": False, "message": str(e)}

    try:
        notifier = TelegramNotifier()
        notifier.send_notification("Test automatico desde Herramientas", "INFO")
        results["telegram"] = {"success": True, "message": "Mensaje enviado"}
    except Exception as e:
        results["telegram"] = {"success": False, "message": str(e)}

    try:
        wp_url = SettingsResolver.get("wp_api_url", "")
        wp_user = SettingsResolver.get("wp_username", "")
        wp_pass = SettingsResolver.get("wp_app_password", "")
        if not wp_url or not wp_user:
            results["wordpress"] = {"success": False, "message": "WordPress no configurado"}
        else:
            import requests as _req
            from requests.auth import HTTPBasicAuth
            resp = _req.get(wp_url.rstrip("/") + "/users/me", auth=HTTPBasicAuth(wp_user, wp_pass), timeout=10)
            if resp.status_code == 200:
                results["wordpress"] = {"success": True, "message": "Conexion OK"}
            else:
                results["wordpress"] = {"success": False, "message": f"HTTP {resp.status_code}"}
    except Exception as e:
        results["wordpress"] = {"success": False, "message": str(e)}

    try:
        SettingsResolver.reload(db)
        results["cache"] = {"success": True, "message": "Cache recargada correctamente"}
    except Exception as e:
        results["cache"] = {"success": False, "message": str(e)}

    total = len(results)
    passed = sum(1 for r in results.values() if r["success"])
    results["_summary"] = {"total": total, "passed": passed, "failed": total - passed}
    return results

@router.post("/test/pipeline-batch")
async def test_pipeline_batch(request: Request):
    import time as _time
    import shutil as _shutil
    import os as _os
    import traceback as _tb

    body = await request.json()
    municipality = body.get("municipality", "")
    category = body.get("category", "")
    do_cleanup = body.get("cleanup", True)

    working_dir = SettingsResolver.get("working_folder_path", "/tmp/editorial_working")
    if not working_dir:
        return {"success": False, "message": "Ruta working_folder_path no configurada"}
    try:
        _os.makedirs(working_dir, exist_ok=True)
    except Exception as e:
        return {"success": False, "message": f"Error creando working dir: {e}"}

    test_folder = _os.path.join(working_dir, f"_test_pipeline_{int(_time.time())}")
    try:
        _os.makedirs(test_folder, exist_ok=True)
        unique_text = f"Contenido de prueba del pipeline. Timestamp: {_time.time()}. Identificador unico para evitar duplicados."

        docx_path = _os.path.join(test_folder, "documento_prueba.docx")
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("Noticia de prueba - Bergueda", level=1)
            doc.add_paragraph(unique_text)
            doc.add_paragraph("El ajuntament de Bergueda ha anunciat noves activitats per a la setmana vinent. La programacio inclou actes culturals, esportius i de medi ambient per a totes les edats.")
            doc.save(docx_path)
            docx_ok = True
        except Exception:
            docx_ok = False

        txt_path = _os.path.join(test_folder, "texto_prueba.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(unique_text + "\n\nNoticia de prueba: Esdeveniment programat per a la propera setmana a la zona del Maresme.")

        files_created = [txt_path]
        if docx_ok:
            files_created.append(docx_path)

        from app.config.settings import settings
        from app.services.pipeline.orchestrator import PipelineOrchestrator
        from app.db.session import SessionLocal as _SessionLocal
        from app.db.repositories.all_repos import source_batch_repo, content_candidate_repo, canonical_content_repo

        pipeline = PipelineOrchestrator(settings)

        batch_id = None
        start = _time.time()

        if municipality:
            tmp_path = test_folder
            pipeline.process_new_batch(tmp_path)
        else:
            pipeline.process_new_batch(test_folder)

        elapsed = round(_time.time() - start, 1)

        db = _SessionLocal()
        try:
            from app.db.models import SourceBatch, ContentCandidate
            batches = db.query(SourceBatch).filter(SourceBatch.original_path == test_folder).order_by(SourceBatch.created_at.desc()).all()
            if not batches:
                batches = db.query(SourceBatch).filter(SourceBatch.external_name.like("_test_pipeline_%")).order_by(SourceBatch.created_at.desc()).limit(1).all()

            if not batches:
                return {
                    "success": True,
                    "message": "Pipeline ejecutado pero no se encontro el lote en la DB (posible duplicado)",
                    "duration": elapsed,
                    "details": {"batch_status": "DUPLICADO_O_MISSING", "files_ingested": len(files_created), "candidates": [], "events": [], "error_trace": None}
                }

            batch = batches[0]
            batch_id = batch.id
            batch_data = {
                "batch_id": str(batch.id),
                "batch_status": batch.status or "UNKNOWN",
                "files_ingested": len(batch.files) if batch.files else 0,
                "candidates": [],
                "events": [],
                "error_trace": batch.error_message or None
            }

            for c in (batch.candidates or []):
                cand_data = {
                    "status": c.status or "UNKNOWN",
                    "municipality": c.municipality,
                    "category": c.category,
                    "confidence": c.classification_confidence,
                    "title": None,
                    "error": c.review_reason
                }
                if c.canonical_content:
                    cc = c.canonical_content
                    cand_data["title"] = cc.final_title
                batch_data["candidates"].append(cand_data)

            if batch_data["candidates"] and batch_data["candidates"][0].get("title"):
                result_msg = f"Pipeline OK ({batch.status}). {len(batch_data['candidates'])} candidatos, {len(files_created)} ficheros"
            elif batch.status == "FAILED":
                result_msg = f"Pipeline fallido: {batch.error_message or 'Error desconocido'}"
            else:
                result_msg = f"Pipeline ejecutado ({batch.status}). {len(batch_data['candidates'])} candidatos generados"

            if do_cleanup and batch_id:
                try:
                    _shutil.rmtree(batch.working_path, ignore_errors=True)
                except Exception:
                    pass
                try:
                    db.delete(batch)
                    db.commit()
                except Exception:
                    db.rollback()

            return {"success": True, "message": result_msg, "duration": elapsed, "details": batch_data}

        finally:
            db.close()

    except Exception as e:
        return {
            "success": False,
            "message": f"Error en el pipeline: {str(e)}",
            "details": {"error_trace": _tb.format_exc()}
        }
    finally:
        if do_cleanup and _os.path.exists(test_folder):
            try:
                _shutil.rmtree(test_folder, ignore_errors=True)
            except Exception:
                pass

@router.get("/{category}", response_class=HTMLResponse)
def settings_section(request: Request, category: str, db: Session = Depends(get_db)):
    """Vista generica para categorias sin plantilla dedicada"""
    items = _get_category_items(db, category)
    return templates.TemplateResponse(
        request=request,
        name=f"settings/form.html",
        context={
            "category": category,
            "items": items,
            "success": request.query_params.get("success")
        }
    )

@router.post("/{category}")
async def save_settings_section(request: Request, category: str, db: Session = Depends(get_db)):
    form_data = await request.form()
    updates = []
    
    # Simple form parsing to Pydantic
    for key, val in form_data.items():
        if key.startswith("setting_"):
            actual_key = key.replace("setting_", "")
            is_secret = form_data.get(f"secret_{actual_key}") == "on"
            val_type = form_data.get(f"type_{actual_key}", SettingType.STRING)
            
            # Type casting
            final_val = val
            if val_type == SettingType.BOOLEAN:
                final_val = True if val.lower() in ['true', 'on', '1', 'yes'] else False
            elif val_type == SettingType.INTEGER:
                final_val = int(val) if val else 0
                
            updates.append(SettingItemUpdate(
                key=actual_key,
                value=final_val,
                value_type=val_type,
                is_secret=is_secret
            ))
            
    SettingsService.update_section(db, category, updates, user="admin_user")
    
    return RedirectResponse(url=f"/settings/{category}?success=true", status_code=303)

# --- Test Endpoints ---

@router.post("/test/telegram")
def test_telegram(db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    bot_token = SettingsResolver.get("telegram_bot_token", "")
    chat_id = SettingsResolver.get("telegram_chat_id", "")
    if not bot_token:
        return {"success": False, "message": "Bot token no configurado. Guarda el token en la seccion de Telegram."}
    if not chat_id:
        return {"success": False, "message": "Chat ID no configurado. Guarda el chat ID en la seccion de Telegram."}
    try:
        import requests as _req
        url = f"https://api.telegram.org/bot{bot_token}/sendMessage"
        resp = _req.post(url, json={
            "chat_id": chat_id,
            "text": "Prueba de conexion desde el Panel de Configuracion. Si ves este mensaje, Telegram funciona correctamente.",
        }, timeout=10)
        data = resp.json()
        if data.get("ok"):
            return {"success": True, "message": f"Mensaje enviado correctamente al chat {chat_id}"}
        else:
            desc = data.get("description", "Error desconocido")
            return {"success": False, "message": f"Error de Telegram: {desc}"}
    except Exception as e:
        return {"success": False, "message": f"Error de conexion: {str(e)}"}

@router.post("/test/ftp")
def test_ftp(db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    client = FtpRemoteInboxClient()
    success, msg = client.test_connection()
    return {"success": success, "message": msg}

@router.post("/test/sftp")
def test_sftp(db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    client = SftpRemoteInboxClient()
    success, msg = client.test_connection()
    return {"success": success, "message": msg}

@router.post("/test/smb")
def test_smb(db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    client = SmbRemoteInboxClient()
    success, msg = client.test_connection()
    return {"success": success, "message": msg}

@router.post("/test/smb/shares")
async def discover_smb_shares(request: Request, db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    client = SmbRemoteInboxClient()
    body = {}
    try:
        body = await request.json()
    except Exception:
        pass
    custom_names = body.get("custom_names", None)
    success, msg, shares = client.discover_shares(custom_names)
    return {"success": success, "message": msg, "shares": shares}

@router.post("/test/smb/folders")
async def list_smb_folders(request: Request, db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    body = await request.json()
    share_name = body.get("share", "")
    client = SmbRemoteInboxClient()
    success, msg, folders = client.list_share_folders(share_name)
    return {"success": success, "message": msg, "folders": folders}

@router.post("/test/smb/write")
def test_smb_write(db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    client = SmbRemoteInboxClient()
    success, msg = client.test_write()
    return {"success": success, "message": msg}

@router.post("/test/smb/write-folder")
async def test_smb_write_folder(request: Request, db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    body = await request.json()
    base_path = body.get("base_path", "/")
    processed_path = body.get("processed_path", "/processed")
    delete_after = body.get("delete_after", True)
    client = SmbRemoteInboxClient()
    success, msg = client.test_write_folder(base_path, processed_path, delete_after)
    return {"success": success, "message": msg}

@router.post("/test/smb/subfolders")
async def list_smb_subfolders(request: Request, db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    body = await request.json()
    base_path = body.get("base_path", "")
    client = SmbRemoteInboxClient()
    success, msg, folders = client.list_subfolders(base_path)
    return {"success": success, "message": msg, "folders": folders}

@router.post("/test/local")
def test_local(db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    client = LocalFolderInboxClient()
    success, msg = client.test_connection()
    return {"success": success, "message": msg}

@router.post("/test/local/write-folder")
async def test_local_write_folder(request: Request, db: Session = Depends(get_db)):
    import os
    import tempfile
    from datetime import datetime
    SettingsResolver.reload(db)
    body = await request.json()
    base_path = body.get("base_path", "/")
    processed_path = body.get("processed_path", "/processed")
    delete_after = body.get("delete_after", True)

    local_base = SettingsResolver.get("hot_folder_local_path", "/tmp/hot_folder")
    if not local_base:
        return {"success": False, "message": "Ruta base local no configurada"}

    clean_base = base_path.lstrip("/")
    full_path = os.path.join(local_base, clean_base) if clean_base else local_base
    clean_processed = processed_path.lstrip("/")
    full_processed = os.path.join(local_base, clean_processed) if clean_processed else os.path.join(local_base, "processed")

    test_filename = "test_connection.txt"
    test_content = f"Local write test - {base_path} - PANXING\n"
    results = []
    label = base_path.strip("/").split("/")[-1] or "raiz"

    if not os.path.exists(full_path):
        return {"success": False, "message": f"La ruta no existe: {full_path}"}
    if not os.path.isdir(full_path):
        return {"success": False, "message": f"No es un directorio: {full_path}"}

    test_file_hot = os.path.join(full_path, test_filename)
    try:
        with open(test_file_hot, "w") as f:
            f.write(test_content)
        if delete_after:
            os.remove(test_file_hot)
            results.append(f"[{label}] Hotfolder: escritura + borrado OK ({full_path})")
        else:
            results.append(f"[{label}] Hotfolder: archivo DEJADO en {full_path}")
    except Exception as e:
        results.append(f"[{label}] Hotfolder FALLO ({full_path}): {e}")
        return {"success": False, "message": " | ".join(results)}

    try:
        os.makedirs(full_processed, exist_ok=True)
        test_file_proc = os.path.join(full_processed, test_filename)
        with open(test_file_proc, "w") as f:
            f.write(test_content)
        if delete_after:
            os.remove(test_file_proc)
            results.append(f"[{label}] Procesados: escritura + borrado OK ({full_processed})")
        else:
            results.append(f"[{label}] Procesados: archivo DEJADO en {full_processed}")
    except Exception as e:
        results.append(f"[{label}] Procesados FALLO ({full_processed}): {e}")
        return {"success": False, "message": " | ".join(results)}

    return {"success": True, "message": " | ".join(results)}

@router.post("/test/outfolder/ftp")
def test_outfolder_ftp(db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    client = FtpOutfolderClient()
    success, msg = client.test_connection()
    return {"success": success, "message": msg}

@router.post("/test/outfolder/local")
async def test_outfolder_local(request: Request, db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    body = await request.json()
    path = body.get("path", "")
    client = LocalOutfolderClient()
    success, msg = client.test_connection(path)
    return {"success": success, "message": msg}

@router.post("/test/outfolder/write-folder")
async def test_outfolder_write_folder(request: Request, db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    body = await request.json()
    base_path = body.get("base_path", "/")
    delete_after = body.get("delete_after", True)
    mode = body.get("mode", "ftp")
    if mode == "local":
        client = LocalOutfolderClient()
    else:
        client = FtpOutfolderClient()
    success, msg = client.test_write_folder(base_path, delete_after)
    return {"success": success, "message": msg}

@router.post("/test/outfolder/subfolders")
async def list_outfolder_subfolders(request: Request, db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    body = await request.json()
    base_path = body.get("base_path", "")
    mode = body.get("mode", "ftp")
    if mode == "local":
        client = LocalOutfolderClient()
    else:
        client = FtpOutfolderClient()
    success, msg, folders = client.list_subfolders(base_path)
    return {"success": success, "message": msg, "folders": folders}

@router.post("/test/llm")
async def test_llm(request: Request, db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    body = {}
    try:
        body = await request.json()
    except Exception:
        pass
    connection_id = body.get("connection_id", "")

    connections_json = SettingsResolver.get("llm_connections", "[]")
    try:
        connections = _json.loads(connections_json)
    except Exception:
        connections = []

    if connection_id:
        conn = next((c for c in connections if c.get("id") == connection_id), None)
    else:
        conn = next((c for c in connections if c.get("active") and c.get("enabled")), None)

    if not conn:
        return {"success": False, "message": "No hay conexion activa o la conexion indicada no existe."}
    provider = conn.get("provider", "openai")
    api_key = conn.get("api_key", "")
    model = conn.get("model", "")
    if provider != "ollama" and not api_key:
        return {"success": False, "message": f"Clave API no configurada para {provider}."}
    try:
        import requests as _req
        if provider == "openai":
            url = "https://api.openai.com/v1/chat/completions"
            headers = {"Authorization": f"Bearer {api_key}"}
            payload = {"model": model or "gpt-4o-mini", "messages": [{"role": "user", "content": "Responde OK"}], "max_tokens": 5}
        elif provider == "anthropic":
            url = "https://api.anthropic.com/v1/messages"
            headers = {"x-api-key": api_key, "anthropic-version": "2023-06-01", "content-type": "application/json"}
            payload = {"model": model or "claude-3-haiku-20240307", "max_tokens": 5, "messages": [{"role": "user", "content": "Responde OK"}]}
        elif provider == "groq":
            url = "https://api.groq.com/openai/v1/chat/completions"
            headers = {"Authorization": f"Bearer {api_key}"}
            payload = {"model": model, "messages": [{"role": "user", "content": "Responde OK"}], "max_tokens": 5}
        elif provider == "mistral":
            url = "https://api.mistral.ai/v1/chat/completions"
            headers = {"Authorization": f"Bearer {api_key}"}
            payload = {"model": model, "messages": [{"role": "user", "content": "Responde OK"}], "max_tokens": 5}
        elif provider == "gemini":
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
            headers = {}
            payload = {"contents": [{"parts": [{"text": "Responde OK"}]}], "generationConfig": {"maxOutputTokens": 5}}
        elif provider == "ollama":
            url = "http://localhost:11434/api/chat"
            headers = {"Content-Type": "application/json"}
            payload = {"model": model, "messages": [{"role": "user", "content": "Responde OK"}], "stream": False}
        elif provider == "azure":
            return {"success": False, "message": "Azure requiere URL custom. Usa el test desde el pipeline."}
        else:
            return {"success": False, "message": f"Proveedor {provider} no soportado."}
        resp = _req.post(url, json=payload, headers=headers, timeout=15)
        resp.raise_for_status()
        return {"success": True, "message": f"Conexion exitosa: {provider} ({model})"}
    except Exception as e:
        return {"success": False, "message": _format_llm_test_error(provider, model, e)}

@router.post("/test/fetch-models")
async def fetch_llm_models(request: Request):
    body = await request.json()
    provider = body.get("provider", "")
    api_key = body.get("api_key", "")
    models = []
    try:
        import requests as _req
        if provider == "openai":
            resp = _req.get("https://api.openai.com/v1/models", headers={"Authorization": f"Bearer {api_key}"}, timeout=15)
            resp.raise_for_status()
            for m in resp.json().get("data", []):
                mid = m.get("id", "")
                if mid.startswith("ft-") or any(mid.startswith(x) for x in ["babbage", "davinci", "curie", "ada.", "text-", "tts-", "dall-e", "whisper", "gpt-3.5-turbo-instruct"]):
                    continue
                if not any(x in mid for x in ["gpt-4", "gpt-3.5", "o1", "o3", "o4", "chatgpt", "gpt-4o", "gpt-4.5"]):
                    continue
                vision = any(x in mid for x in ["4o", "gpt-4-turbo", "gpt-4-1106", "gpt-4.5"])
                models.append({"id": mid, "name": mid, "vision": vision})
            models.sort(key=lambda x: x["id"])
        elif provider == "anthropic":
            try:
                resp = _req.get("https://api.anthropic.com/v1/models", headers={"x-api-key": api_key, "anthropic-version": "2023-06-01"}, timeout=15)
                resp.raise_for_status()
                for m in resp.json().get("data", []):
                    mid = m.get("id", "")
                    models.append({"id": mid, "name": m.get("display_name", mid), "vision": True})
                models.sort(key=lambda x: x["id"])
            except Exception:
                models = [
                    {"id": "claude-sonnet-4-20250514", "name": "Claude Sonnet 4", "vision": True},
                    {"id": "claude-3-5-sonnet-20241022", "name": "Claude 3.5 Sonnet", "vision": True},
                    {"id": "claude-3-5-haiku-20241022", "name": "Claude 3.5 Haiku", "vision": True},
                    {"id": "claude-3-opus-20240229", "name": "Claude 3 Opus", "vision": True},
                    {"id": "claude-3-sonnet-20240229", "name": "Claude 3 Sonnet", "vision": True},
                    {"id": "claude-3-haiku-20240307", "name": "Claude 3 Haiku", "vision": True},
                ]
        elif provider == "groq":
            resp = _req.get("https://api.groq.com/openai/v1/models", headers={"Authorization": f"Bearer {api_key}"}, timeout=15)
            resp.raise_for_status()
            for m in resp.json().get("data", []):
                mid = m.get("id", "")
                vision = "vision" in mid.lower()
                models.append({"id": mid, "name": mid, "vision": vision})
            models.sort(key=lambda x: x["id"])
        elif provider == "mistral":
            resp = _req.get("https://api.mistral.ai/v1/models", headers={"Authorization": f"Bearer {api_key}"}, timeout=15)
            resp.raise_for_status()
            for m in resp.json().get("data", []):
                mid = m.get("id", "")
                vision = "pixtral" in mid.lower()
                models.append({"id": mid, "name": mid, "vision": vision})
            models.sort(key=lambda x: x["id"])
        elif provider == "gemini":
            resp = _req.get(f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key}", timeout=15)
            resp.raise_for_status()
            for m in resp.json().get("models", []):
                mid = m.get("name", "").replace("models/", "")
                methods = m.get("supportedGenerationMethods", [])
                vision = "generateContent" in methods and "pro" in mid.lower()
                models.append({"id": mid, "name": mid, "vision": vision})
            models.sort(key=lambda x: x["id"])
        elif provider == "ollama":
            resp = _req.get("http://localhost:11434/api/tags", timeout=5)
            resp.raise_for_status()
            for m in resp.json().get("models", []):
                mid = m.get("name", "")
                vision = any(x in mid.lower() for x in ["llava", "bakllava", "moondream", "minicpm-v", "llama3.2-vision", "pixtral"])
                models.append({"id": mid, "name": mid, "vision": vision})
            models.sort(key=lambda x: x["id"])
        elif provider == "azure":
            return {"success": False, "message": "Azure no permite listar modelos automaticamente. Introduce el nombre del deployment manualmente.", "models": []}
        else:
            return {"success": False, "message": f"Proveedor {provider} no soportado.", "models": []}
        return {"success": True, "message": f"{len(models)} modelos encontrados.", "models": models}
    except Exception as e:
        provider_name = provider or "LLM"
        if provider == "gemini":
            return {
                "success": False,
                "message": (
                    "No se pudieron cargar los modelos de Gemini. "
                    f"{_format_llm_test_error(provider, '', e)}"
                ),
                "models": [],
            }
        return {"success": False, "message": f"No se pudieron cargar los modelos de {provider_name}: {_extract_http_error_message(e)}", "models": []}

@router.post("/test/wordpress")
async def test_wordpress(request: Request, db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    api_url = SettingsResolver.get("wp_api_url", "")
    username = SettingsResolver.get("wp_username", "")
    app_password = SettingsResolver.get("wp_app_password", "")
    if not api_url:
        return {"success": False, "message": "URL de la API de WordPress no configurada."}
    if not username or not app_password:
        return {"success": False, "message": "Usuario o contraseña de aplicacion no configurados."}
    try:
        import requests
        from requests.auth import HTTPBasicAuth
        test_url = api_url.rstrip("/") + "/users/me"
        resp = requests.get(test_url, auth=HTTPBasicAuth(username, app_password), timeout=10)
        if resp.status_code == 200:
            user_data = resp.json()
            return {"success": True, "message": f"Conexion exitosa. Usuario: {user_data.get('name', username)} (ID: {user_data.get('id', '?')})"}
        else:
            return {"success": False, "message": f"Error HTTP {resp.status_code}: {resp.text[:200]}"}
    except Exception as e:
        return {"success": False, "message": str(e)}

@router.post("/test/path")
async def test_path(request: Request):
    body = await request.json()
    path = body.get("path", "")
    import os
    if not path:
        return {"success": False, "message": "Ruta vacia"}
    if os.path.exists(path):
        writable = os.access(path, os.W_OK)
        return {"success": True, "writable": writable, "message": "Existe" + ("" if writable else " pero sin permisos de escritura")}
    else:
        return {"success": False, "writable": False, "message": "La ruta no existe"}

@router.post("/test/path-create")
async def test_path_create(request: Request):
    body = await request.json()
    paths = body.get("paths", [])
    import os
    created = []
    failed = []
    for p in paths:
        if not p:
            continue
        try:
            os.makedirs(p, exist_ok=True)
            created.append(p)
        except Exception as e:
            failed.append(f"{p}: {e}")
    if failed:
        return {"success": False, "message": f"Creadas: {len(created)}. Errores: {'; '.join(failed)}"}
    return {"success": True, "message": f"{len(created)} carpetas verificadas/creadas correctamente."}

@router.post("/reload")
def reload_config(db: Session = Depends(get_db)):
    SettingsResolver.reload(db)
    return {"success": True, "message": "Configuración recargada en memoria"}
