import re
import os
import logging
from typing import List, Dict, Any, Optional
from abc import ABC, abstractmethod
from uuid import UUID

from app.schemas.all_schemas import ExtractionResult
from app.core.enums import ExtractionMethod
from app.services.extraction.cleaning import TextCleaningPipeline
from app.services.settings.service import SettingsResolver

OCR_ENGINE = None
logger = logging.getLogger(__name__)


class BaseExtractor(ABC):
    def __init__(self):
        self.cleaner = TextCleaningPipeline()

    @abstractmethod
    def extract(self, file_path: str, file_id: str) -> ExtractionResult:
        pass

    def _uuid(self, file_id: str) -> UUID:
        if isinstance(file_id, UUID):
            return file_id
        return UUID(file_id)


class RealPdfExtractor(BaseExtractor):
    def extract(self, file_path: str, file_id: str) -> ExtractionResult:
        try:
            import pymupdf
            doc = pymupdf.open(file_path)
            pages_text = []
            for page in doc:
                text = page.get_text("text")
                if text.strip():
                    pages_text.append(text)
            doc.close()
            raw_text = "\n\n".join(pages_text) if pages_text else "[PDF sin texto extraible]"
        except Exception as e:
            raw_text = f"[Error extrayendo PDF: {e}]"

        cleaned_info = self.cleaner.clean(raw_text)

        return ExtractionResult(
            source_file_id=self._uuid(file_id),
            method=ExtractionMethod.NATIVE_PDF,
            confidence=0.95 + cleaned_info["adjustment"],
            raw_text=raw_text,
            cleaned_text=cleaned_info["cleaned_text"]
        )


class RealDocxExtractor(BaseExtractor):
    def extract(self, file_path: str, file_id: str) -> ExtractionResult:
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    tables_text.append(" | ".join(cells))

            raw_text = "\n\n".join(paragraphs)
            if tables_text:
                raw_text += "\n\n[ TABLAS ]\n" + "\n".join(tables_text)
            if not raw_text.strip():
                raw_text = "[DOCX sin contenido de texto]"
        except Exception as e:
            raw_text = f"[Error extrayendo DOCX: {e}]"

        cleaned_info = self.cleaner.clean(raw_text)

        return ExtractionResult(
            source_file_id=self._uuid(file_id),
            method=ExtractionMethod.DOCX_PARSER,
            confidence=0.98 + cleaned_info["adjustment"],
            raw_text=raw_text,
            cleaned_text=cleaned_info["cleaned_text"]
        )


class RealTextExtractor(BaseExtractor):
    def extract(self, file_path: str, file_id: str) -> ExtractionResult:
        raw_text = ""
        extension = os.path.splitext(file_path)[1].lower()
        for encoding in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    raw_text = f.read()
                break
            except UnicodeDecodeError:
                continue
            except Exception as e:
                raw_text = f"[Error extrayendo texto: {e}]"
                break

        if not str(raw_text).strip():
            raw_text = "[Archivo de texto sin contenido]"

        normalized_text = self._normalize_markdown(raw_text) if extension in (".md", ".markdown") else raw_text

        cleaned_info = self.cleaner.clean(normalized_text)

        return ExtractionResult(
            source_file_id=self._uuid(file_id),
            method=ExtractionMethod.TEXT_FILE,
            confidence=0.96 + cleaned_info["adjustment"],
            raw_text=raw_text,
            cleaned_text=cleaned_info["cleaned_text"]
        )

    def _normalize_markdown(self, text: str) -> str:
        normalized = str(text or "").replace("\r\n", "\n").replace("\r", "\n")

        normalized = re.sub(r"\A---\n.*?\n---\n", "", normalized, flags=re.DOTALL)
        normalized = re.sub(r"^(#{1,6})\s+(.+)$", lambda m: self._markdown_heading_to_marker(m.group(1), m.group(2)), normalized, flags=re.MULTILINE)
        normalized = re.sub(r"```[\s\S]*?```", lambda match: self._strip_code_fence(match.group(0)), normalized)
        normalized = re.sub(r"`([^`]+)`", r"\1", normalized)
        normalized = re.sub(r"!\[([^\]]*)\]\([^\)]+\)", r"\1", normalized)
        normalized = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", normalized)
        normalized = re.sub(r"^>\s?", "", normalized, flags=re.MULTILINE)
        normalized = re.sub(r"^\s*[-*+]\s+(.+)$", r"\n\n[[LI]] \1\n", normalized, flags=re.MULTILINE)
        normalized = re.sub(r"^\s*\d+\.\s+(.+)$", r"\n\n[[LI]] \1\n", normalized, flags=re.MULTILINE)
        normalized = re.sub(r"^(?:---|\*\*\*|___)\s*$", "", normalized, flags=re.MULTILINE)
        normalized = re.sub(r"\*\*([^*]+)\*\*", r"\1", normalized)
        normalized = re.sub(r"__([^_]+)__", r"\1", normalized)
        normalized = re.sub(r"\*([^*]+)\*", r"\1", normalized)
        normalized = re.sub(r"_([^_]+)_", r"\1", normalized)
        normalized = re.sub(r"~~([^~]+)~~", r"\1", normalized)
        normalized = re.sub(r"\|", " ", normalized)
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        return normalized.strip()

    def _markdown_heading_to_marker(self, hashes: str, text: str) -> str:
        level = min(len(hashes), 6)
        clean_text = re.sub(r"\s+", " ", str(text or "")).strip()
        return f"\n[[H{level}]] {clean_text}\n"

    def _strip_code_fence(self, fenced_block: str) -> str:
        lines = fenced_block.splitlines()
        if len(lines) >= 2:
            return "\n".join(lines[1:-1])
        return fenced_block


class ImageOcrExtractor(BaseExtractor):
    def extract(self, file_path: str, file_id: str) -> ExtractionResult:
        global OCR_ENGINE
        engine = SettingsResolver.get("ocr_engine", "disabled")
        if not engine or engine == "disabled":
            return self._mock_extract(file_path, file_id)

        if engine == "ai_vision":
            return self._ai_vision_extract(file_path, file_id)
        elif engine == "tesseract":
            return self._tesseract_extract(file_path, file_id)
        elif engine == "paddleocr":
            return self._paddleocr_extract(file_path, file_id)
        else:
            return self._mock_extract(file_path, file_id)

    def _ai_vision_extract(self, file_path: str, file_id: str) -> ExtractionResult:
        try:
            import base64
            from app.services.ai.client import get_active_llm_client
            from app.services.settings.service import SettingsResolver

            selected_connection_id = str(SettingsResolver.get("ocr_vision_connection_id") or "").strip()
            if not selected_connection_id:
                return self._mock_extract_with_error(file_path, file_id, "ocr_vision_connection_id no configurado")

            client = get_active_llm_client(use_ocr_vision=True)
            if not client:
                return self._mock_extract_with_error(file_path, file_id, "conexion de OCR vision no disponible o no valida")

            with open(file_path, "rb") as f:
                img_b64 = base64.b64encode(f.read()).decode("utf-8")

            prompt = "Extrae todo el texto visible de esta imagen. Responde SOLO con el texto, sin explicaciones."
            response = client.chat(prompt, images=[{"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"}}])

            raw_text = response if response else "[Sin respuesta del modelo]"
            if not str(raw_text).strip() or str(raw_text).strip() == "[Sin respuesta del modelo]":
                return self._empty_ocr_result(file_path, file_id, "Sin respuesta del modelo")
            cleaned_info = self.cleaner.clean(raw_text)

            return ExtractionResult(
                source_file_id=self._uuid(file_id),
                method=ExtractionMethod.OCR_IMAGE,
                confidence=0.85 + cleaned_info["adjustment"],
                raw_text=raw_text,
                cleaned_text=cleaned_info["cleaned_text"]
            )
        except ModuleNotFoundError as e:
            logger.warning("Dependencia OCR/vision no disponible: %s", e)
            return self._mock_extract_with_error(file_path, file_id, "dependencia OCR vision no disponible")
        except Exception as e:
            return self._mock_extract_with_error(file_path, file_id, str(e))

    def _tesseract_extract(self, file_path: str, file_id: str) -> ExtractionResult:
        try:
            import pytesseract
            from PIL import Image as PILImage

            img = PILImage.open(file_path)
            text = pytesseract.image_to_string(img, lang='spa+cat')
            if not text.strip():
                return self._empty_ocr_result(file_path, file_id, "Imagen sin texto reconocible")

            cleaned_info = self.cleaner.clean(text)

            return ExtractionResult(
                source_file_id=self._uuid(file_id),
                method=ExtractionMethod.OCR_IMAGE,
                confidence=0.75 + cleaned_info["adjustment"],
                raw_text=text,
                cleaned_text=cleaned_info["cleaned_text"]
            )
        except ModuleNotFoundError as e:
            logger.warning("Tesseract no disponible: %s", e)
            return self._mock_extract_with_error(file_path, file_id, "pytesseract no instalado en venv")
        except Exception as e:
            return self._mock_extract_with_error(file_path, file_id, str(e))

    def _paddleocr_extract(self, file_path: str, file_id: str) -> ExtractionResult:
        try:
            from paddleocr import PaddleOCR
            ocr = PaddleOCR(use_textline_orientation=True, lang='es')
            result = ocr.ocr(file_path)
            lines = []
            scores = []
            if result and result[0]:
                first = result[0]
                if isinstance(first, dict):
                    rec_texts = first.get("rec_texts") or []
                    rec_scores = first.get("rec_scores") or []
                    for item in rec_texts:
                        text_item = str(item or "").strip()
                        if text_item:
                            lines.append(text_item)
                    for score in rec_scores:
                        try:
                            scores.append(float(score))
                        except Exception:
                            continue
                elif isinstance(first, list):
                    for line in first:
                        if line and len(line) >= 2 and isinstance(line[1], (list, tuple)) and line[1]:
                            text_item = str(line[1][0] or "").strip()
                            if text_item:
                                lines.append(text_item)
                            if len(line[1]) > 1:
                                try:
                                    scores.append(float(line[1][1]))
                                except Exception:
                                    continue
            if not lines:
                return self._empty_ocr_result(file_path, file_id, "Imagen sin texto reconocible")
            text = "\n".join(lines)
            cleaned_info = self.cleaner.clean(text)
            confidence = 0.75 + cleaned_info["adjustment"]
            if scores:
                confidence = max(0.0, min(1.0, sum(scores) / len(scores)))
            return ExtractionResult(
                source_file_id=self._uuid(file_id),
                method=ExtractionMethod.OCR_IMAGE,
                confidence=confidence,
                raw_text=text,
                cleaned_text=cleaned_info["cleaned_text"]
            )
        except ModuleNotFoundError as e:
            logger.warning("PaddleOCR no disponible: %s", e)
            return self._mock_extract_with_error(file_path, file_id, "paddleocr no instalado en venv")
        except Exception as e:
            return self._mock_extract_with_error(file_path, file_id, str(e))

    def _mock_extract(self, file_path: str, file_id: str) -> ExtractionResult:
        fname = os.path.basename(file_path)
        raw_text = f"[OCR deshabilitado para {fname}. Texto no extraido.]"
        return ExtractionResult(
            source_file_id=self._uuid(file_id),
            method=ExtractionMethod.OCR_IMAGE,
            confidence=0.0,
            raw_text=raw_text,
            cleaned_text=""
        )

    def _mock_extract_with_error(self, file_path: str, file_id: str, error: str) -> ExtractionResult:
        raw_text = f"[Error OCR en {os.path.basename(file_path)}: {error}]"
        return ExtractionResult(
            source_file_id=self._uuid(file_id),
            method=ExtractionMethod.OCR_IMAGE,
            confidence=0.0,
            raw_text=raw_text,
            cleaned_text=""
        )

    def _empty_ocr_result(self, file_path: str, file_id: str, reason: str) -> ExtractionResult:
        raw_text = f"[{reason} en {os.path.basename(file_path)}]"
        return ExtractionResult(
            source_file_id=self._uuid(file_id),
            method=ExtractionMethod.OCR_IMAGE,
            confidence=0.0,
            raw_text=raw_text,
            cleaned_text=""
        )


class ExtractionOrchestrator:
    def __init__(self):
        self.pdf_extractor = RealPdfExtractor()
        self.docx_extractor = RealDocxExtractor()
        self.text_extractor = RealTextExtractor()
        self.image_extractor = ImageOcrExtractor()

    def process_files(self, files_info: List[Dict[str, Any]]) -> List[ExtractionResult]:
        results = []
        for file_info in files_info:
            path = file_info["path"].lower()
            file_id = file_info["id"]

            if path.endswith(".pdf"):
                results.append(self.pdf_extractor.extract(path, file_id))
            elif path.endswith(".docx"):
                results.append(self.docx_extractor.extract(path, file_id))
            elif path.endswith((".md", ".markdown", ".txt")):
                results.append(self.text_extractor.extract(path, file_id))
            elif path.endswith((".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp")):
                results.append(self.image_extractor.extract(path, file_id))

        return results
